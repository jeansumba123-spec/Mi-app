from __future__ import annotations

import io
from datetime import datetime
from typing import List, Optional

import numpy as np
import pandas as pd
import plotly.express as px
import plotly.io as pio
import streamlit as st
from scipy.stats import chi2_contingency

try:
    from docx import Document
    from docx.shared import Inches
except Exception:  # pragma: no cover
    Document = None
    Inches = None

st.set_page_config(
    page_title="Análisis de encuestas de movilidad",
    page_icon="🌈",
    layout="wide",
)

PALETTE = {
    "primary": "#1A548C",
    "primary_dark": "#0F304F",
    "accent": "#008C7A",
    "accent_2": "#7C3AED",
    "accent_3": "#F97316",
    "accent_4": "#EC4899",
    "soft": "#F0F7FA",
    "soft_2": "#EEF2FF",
    "danger": "#BF332E",
    "success": "#338C40",
    "warning": "#ED9E1F",
    "white": "#FFFFFF",
    "text": "#152033",
}

CHART_COLORS = [
    "#1A548C",
    "#008C7A",
    "#7C3AED",
    "#F97316",
    "#EC4899",
    "#22C55E",
    "#0EA5E9",
    "#F59E0B",
    "#EF4444",
    "#14B8A6",
]

VARIABLE_CATALOG = {
    "Demográficas": [
        ("Edad", ["edad"]),
        ("Género al nacer", ["sexo", "genero", "género"]),
        ("Ingresos", ["ingreso", "renta"]),
        ("Nivel de instrucción", ["nivel de instruccion", "nivel de instrucción", "instruccion", "instrucción"]),
        ("Zona de residencia", ["zona de residencia", "residencia"]),
    ],
    "Uso": [
        ("Estación", ["estacion", "estación"]),
        ("Zona de destino", ["zona de destino", "destino"]),
        ("Motivo de viaje", ["motivo de viaje", "motivo"]),
        ("Frecuencia de uso", ["frecuencia", "utiliza"]),
        ("Transporte anterior", ["transporte anterior"]),
        ("Bus anterior", ["bus anterior"]),
    ],
    "Percepción": [
        ("Satisfacción", ["satisfaccion", "satisfacción", "grado de satisfaccion", "grado de satisfacción"]),
        ("Factor de cambio", ["factor de cambio", "factor"]),
        ("Sugerencia", ["sugerencia", "observacion", "observación"]),
    ],
}

BIVARIABLE_CATALOG = [
    ("Género vs Satisfacción", ["sexo", "genero", "género"], ["satisfaccion", "satisfacción"]),
    ("Edad vs Frecuencia de uso", ["edad"], ["frecuencia", "utiliza"]),
    ("Ingresos vs Transporte anterior", ["ingreso", "renta"], ["transporte anterior"]),
    ("Motivo de viaje vs Frecuencia de uso", ["motivo de viaje", "motivo"], ["frecuencia", "utiliza"]),
    ("Nivel de instrucción vs Satisfacción", ["nivel de instruccion", "nivel de instrucción", "instruccion", "instrucción"], ["satisfaccion", "satisfacción"]),
    ("Edad vs Factor de cambio", ["edad"], ["factor de cambio", "factor"]),
    ("Ingresos vs Factor de cambio", ["ingreso", "renta"], ["factor de cambio", "factor"]),
    ("Estación vs Motivo de viaje", ["estacion", "estación"], ["motivo de viaje", "motivo"]),
]


def init_state() -> None:
    defaults = {
        "workbooks": {},
        "active_file": None,
        "history": [],
        "records": [],
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


def inject_css() -> None:
    st.markdown(
        f"""
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700;800&display=swap');

        html, body, [class*="css"] {{
            font-family: 'Inter', sans-serif;
        }}

        .stApp {{
            background:
                radial-gradient(circle at top left, rgba(124, 58, 237, 0.15), transparent 32%),
                radial-gradient(circle at top right, rgba(249, 115, 22, 0.16), transparent 28%),
                linear-gradient(135deg, #F8FBFF 0%, #EEF7FF 48%, #FFF7ED 100%);
            color: {PALETTE['text']};
        }}

        section[data-testid="stSidebar"] {{
            background: linear-gradient(180deg, {PALETTE['primary_dark']} 0%, #143B63 45%, #0F766E 100%);
            border-right: 1px solid rgba(255,255,255,.18);
        }}

        section[data-testid="stSidebar"] h1,
        section[data-testid="stSidebar"] h2,
        section[data-testid="stSidebar"] h3,
        section[data-testid="stSidebar"] label,
        section[data-testid="stSidebar"] p,
        section[data-testid="stSidebar"] span,
        section[data-testid="stSidebar"] div {{
            color: white !important;
        }}

        .hero-card {{
            background: linear-gradient(135deg, {PALETTE['primary']} 0%, {PALETTE['accent_2']} 55%, {PALETTE['accent_4']} 100%);
            padding: 28px 30px;
            border-radius: 26px;
            color: white;
            box-shadow: 0 18px 45px rgba(26, 84, 140, .24);
            margin-bottom: 20px;
        }}

        .hero-title {{
            font-size: 34px;
            line-height: 1.1;
            font-weight: 800;
            margin: 0 0 8px 0;
        }}

        .hero-subtitle {{
            font-size: 16px;
            opacity: .93;
            margin: 0;
        }}

        .color-title {{
            background: linear-gradient(90deg, {PALETTE['primary']}, {PALETTE['accent_2']}, {PALETTE['accent_3']});
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            font-size: 31px;
            font-weight: 800;
            margin-bottom: 10px;
        }}

        .section-card {{
            background: rgba(255,255,255,.82);
            border: 1px solid rgba(15, 48, 79, .08);
            border-radius: 22px;
            padding: 20px;
            box-shadow: 0 14px 30px rgba(15, 48, 79, .08);
            margin: 12px 0 18px 0;
        }}

        .mini-card {{
            padding: 18px;
            border-radius: 20px;
            color: white;
            min-height: 118px;
            box-shadow: 0 12px 25px rgba(15, 48, 79, .13);
        }}

        .mini-card h3 {{
            margin: 0 0 7px 0;
            font-size: 22px;
            font-weight: 800;
            color: white;
        }}

        .mini-card p {{
            margin: 0;
            color: white;
            opacity: .93;
        }}

        .metric-card {{
            border-radius: 18px;
            padding: 16px 18px;
            background: white;
            border-left: 7px solid {PALETTE['accent_2']};
            box-shadow: 0 10px 25px rgba(15, 48, 79, .08);
        }}

        .metric-label {{
            font-size: 13px;
            color: #64748B;
            font-weight: 700;
            margin-bottom: 6px;
        }}

        .metric-value {{
            font-size: 25px;
            font-weight: 800;
            color: {PALETTE['primary_dark']};
        }}

        div[data-testid="stButton"] > button {{
            border-radius: 14px;
            border: 0;
            font-weight: 800;
            color: white;
            background: linear-gradient(90deg, {PALETTE['accent']} 0%, {PALETTE['accent_2']} 100%);
            box-shadow: 0 9px 18px rgba(0, 140, 122, .20);
        }}

        div[data-testid="stDownloadButton"] > button {{
            border-radius: 14px;
            border: 0;
            font-weight: 800;
            color: white;
            background: linear-gradient(90deg, {PALETTE['primary']} 0%, {PALETTE['accent_4']} 100%);
        }}

        .history-pill {{
            display: inline-block;
            background: linear-gradient(90deg, {PALETTE['accent_2']}, {PALETTE['accent_4']});
            color: white;
            padding: 6px 12px;
            border-radius: 999px;
            font-size: 13px;
            font-weight: 800;
            margin-bottom: 8px;
        }}
        </style>
        """,
        unsafe_allow_html=True,
    )


def page_header(title: str, subtitle: str, icon: str = "🌈") -> None:
    st.markdown(
        f"""
        <div class="hero-card">
            <div class="hero-title">{icon} {title}</div>
            <p class="hero-subtitle">{subtitle}</p>
        </div>
        """,
        unsafe_allow_html=True,
    )


def section_title(title: str) -> None:
    st.markdown(f'<div class="color-title">{title}</div>', unsafe_allow_html=True)


def metric_card(label: str, value: str, accent: str = "#7C3AED") -> None:
    st.markdown(
        f"""
        <div class="metric-card" style="border-left-color:{accent};">
            <div class="metric-label">{label}</div>
            <div class="metric-value">{value}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def add_history(text: str, record: Optional[dict] = None) -> None:
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    st.session_state.history.append(f"{timestamp} | {text}")
    if record is not None:
        record["timestamp"] = timestamp
        record["summary_text"] = text
        st.session_state.records.append(record)


def normalize_name(value: str) -> str:
    return str(value).strip().lower()


def find_column(data: pd.DataFrame, patterns: List[str]) -> Optional[str]:
    columns = list(data.columns)
    for pattern in patterns:
        p = normalize_name(pattern)
        for col in columns:
            if p in normalize_name(col):
                return col
    return None


def clean_series(series: pd.Series) -> pd.Series:
    cleaned = series.dropna()
    if cleaned.dtype == object:
        cleaned = cleaned.astype(str).str.strip()
        cleaned = cleaned[cleaned != ""]
    return cleaned


def as_labels(series: pd.Series) -> pd.Series:
    return clean_series(series).astype(str).str.strip()


def as_numeric(series: pd.Series) -> pd.Series:
    return pd.to_numeric(clean_series(series), errors="coerce").dropna()


def frequency_table(series: pd.Series) -> pd.DataFrame:
    labels = as_labels(series)
    counts = labels.value_counts(dropna=True, sort=False)
    result = pd.DataFrame({"Categoría": counts.index.astype(str), "Frecuencia": counts.values})
    total = result["Frecuencia"].sum()
    result["Porcentaje"] = np.where(total > 0, result["Frecuencia"] / total * 100, 0).round(2)
    return result


def load_workbook(uploaded_file) -> pd.DataFrame:
    name = uploaded_file.name.lower()
    if name.endswith(".csv"):
        return pd.read_csv(uploaded_file)
    if name.endswith((".xlsx", ".xls")):
        return pd.read_excel(uploaded_file)
    raise ValueError("Formato no soportado. Usa CSV, XLSX o XLS.")


def style_figure(fig, title: str):
    fig.update_layout(
        title={"text": title, "x": 0.02, "xanchor": "left", "font": {"size": 22, "color": PALETTE["primary_dark"]}},
        paper_bgcolor="rgba(255,255,255,0)",
        plot_bgcolor="rgba(255,255,255,0.90)",
        font={"family": "Inter, Arial", "color": PALETTE["text"]},
        legend_title_text="",
        margin=dict(l=35, r=25, t=65, b=35),
    )
    fig.update_xaxes(showgrid=True, gridcolor="rgba(15,48,79,.08)")
    fig.update_yaxes(showgrid=True, gridcolor="rgba(15,48,79,.08)")
    return fig


def fig_to_png_bytes(fig) -> Optional[bytes]:
    try:
        return fig.to_image(format="png", width=1100, height=650, scale=2)
    except Exception:
        return None


def fig_from_record(record: dict):
    figure_json = record.get("figure_json")
    if not figure_json:
        return None
    try:
        return pio.from_json(figure_json)
    except Exception:
        return None


def sidebar() -> str:
    st.sidebar.title("🌈 Menú")
    st.sidebar.caption("Panel web de análisis de movilidad")

    uploaded = st.sidebar.file_uploader(
        "📂 Cargar archivos Excel o CSV",
        type=["xlsx", "xls", "csv"],
        accept_multiple_files=True,
    )
    if uploaded:
        for file in uploaded:
            try:
                st.session_state.workbooks[file.name] = load_workbook(file)
                st.session_state.active_file = file.name
                st.sidebar.success(f"Cargado: {file.name}")
            except Exception as exc:
                st.sidebar.error(f"No se pudo cargar {file.name}: {exc}")

    if st.session_state.workbooks:
        names = list(st.session_state.workbooks.keys())
        current_index = names.index(st.session_state.active_file) if st.session_state.active_file in names else 0
        st.session_state.active_file = st.sidebar.selectbox("📌 Archivo activo", names, index=current_index)
    else:
        st.sidebar.info("Aún no hay archivos cargados.")

    total_rows = sum(len(df) for df in st.session_state.workbooks.values())
    active_cols = 0
    if st.session_state.active_file:
        active_cols = st.session_state.workbooks[st.session_state.active_file].shape[1]

    st.sidebar.markdown("---")
    st.sidebar.metric("Archivos", len(st.session_state.workbooks))
    st.sidebar.metric("Registros totales", total_rows)
    st.sidebar.metric("Variables activas", active_cols)

    st.sidebar.markdown("---")
    return st.sidebar.radio(
        "Vista",
        ["Inicio", "Análisis univariable", "Análisis bivariable", "Historial y exportación", "Ayuda"],
    )


def active_data() -> Optional[pd.DataFrame]:
    name = st.session_state.active_file
    if not name:
        return None
    return st.session_state.workbooks.get(name)


def home_view() -> None:
    page_header(
        "Análisis de encuestas de movilidad",
        "Carga tus datos, genera tablas, gráficos, cruces estadísticos y guarda todo en el historial.",
        "🚍",
    )

    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown('<div class="mini-card" style="background:linear-gradient(135deg,#1A548C,#0EA5E9);"><h3>📂 Datos</h3><p>Carga uno o varios archivos Excel o CSV.</p></div>', unsafe_allow_html=True)
    with c2:
        st.markdown('<div class="mini-card" style="background:linear-gradient(135deg,#008C7A,#22C55E);"><h3>📊 Análisis</h3><p>Obtén frecuencias, porcentajes y cruces.</p></div>', unsafe_allow_html=True)
    with c3:
        st.markdown('<div class="mini-card" style="background:linear-gradient(135deg,#7C3AED,#EC4899);"><h3>🖼️ Historial</h3><p>Guarda tablas, métricas y gráficos completos.</p></div>', unsafe_allow_html=True)

    data = active_data()
    if data is not None:
        section_title(f"Vista previa: {st.session_state.active_file}")
        st.dataframe(data.head(30), use_container_width=True)
    else:
        st.info("Carga un archivo desde el menú lateral para comenzar.")


def build_univariable_chart(table: pd.DataFrame, label: str, chart_type: str):
    if chart_type == "Barras":
        fig = px.bar(
            table,
            x="Categoría",
            y="Frecuencia",
            title=f"Distribución de {label}",
            text="Frecuencia",
            color="Categoría",
            color_discrete_sequence=CHART_COLORS,
        )
        fig.update_traces(textposition="outside")
    elif chart_type == "Pastel":
        fig = px.pie(
            table,
            names="Categoría",
            values="Frecuencia",
            title=f"Participación porcentual de {label}",
            color_discrete_sequence=CHART_COLORS,
            hole=0.36,
        )
        fig.update_traces(textinfo="percent+label")
    else:
        pareto = table.sort_values("Frecuencia", ascending=False).copy()
        pareto["Porcentaje acumulado"] = pareto["Frecuencia"].cumsum() / pareto["Frecuencia"].sum() * 100
        fig = px.bar(
            pareto,
            x="Categoría",
            y="Frecuencia",
            title=f"Pareto de {label}",
            text="Frecuencia",
            color="Categoría",
            color_discrete_sequence=CHART_COLORS,
        )
        fig.add_scatter(
            x=pareto["Categoría"],
            y=pareto["Porcentaje acumulado"],
            mode="lines+markers",
            name="Porcentaje acumulado",
            yaxis="y2",
            line=dict(width=4),
        )
        fig.update_layout(
            yaxis2=dict(title="Porcentaje acumulado", overlaying="y", side="right", range=[0, 105]),
        )
    return style_figure(fig, fig.layout.title.text)


def univariable_view() -> None:
    page_header("Análisis univariable", "Explora una variable con frecuencias, porcentajes, estadísticos y gráficos coloridos.", "📊")
    data = active_data()
    if data is None:
        st.info("Carga un archivo Excel o CSV desde el menú lateral.")
        return

    st.markdown(f'<span class="history-pill">Archivo activo: {st.session_state.active_file}</span>', unsafe_allow_html=True)

    col_a, col_b, col_c = st.columns([1.1, 1.1, 1])
    with col_a:
        mode = st.radio("Selección de variable", ["Catálogo automático", "Columna manual"], horizontal=True)
    with col_b:
        chart_type = st.selectbox("Tipo de gráfico", ["Barras", "Pastel", "Pareto"])

    if mode == "Catálogo automático":
        with col_c:
            group = st.selectbox("Grupo", list(VARIABLE_CATALOG.keys()))
        label = st.selectbox("Variable", [x[0] for x in VARIABLE_CATALOG[group]])
        patterns = dict(VARIABLE_CATALOG[group])[label]
        column = find_column(data, patterns)
        if column is None:
            st.error(f"No se encontró una columna compatible con: {label}")
            return
    else:
        column = st.selectbox("Columna", data.columns)
        label = column

    table = frequency_table(data[column])
    numeric = as_numeric(data[column])

    section_title(label)
    total_valid = int(table["Frecuencia"].sum()) if not table.empty else 0
    mode_row = table.loc[table["Frecuencia"].idxmax()] if not table.empty else None

    m1, m2, m3, m4 = st.columns(4)
    with m1:
        metric_card("N válido", f"{total_valid:,}", "#1A548C")
    with m2:
        metric_card("Moda", str(mode_row["Categoría"]) if mode_row is not None else "-", "#008C7A")
    with m3:
        metric_card("% moda", f"{mode_row['Porcentaje']:.1f}%" if mode_row is not None else "-", "#F97316")
    with m4:
        metric_card("Categorías", f"{len(table):,}", "#EC4899")

    stats_lines = [f"N válido: {total_valid}"]
    if mode_row is not None:
        stats_lines.append(f"Moda: {mode_row['Categoría']} ({int(mode_row['Frecuencia'])} casos, {mode_row['Porcentaje']:.1f}%)")
    if len(numeric) > 0:
        stats_lines.extend([
            f"Media: {numeric.mean():.2f}",
            f"Mediana: {numeric.median():.2f}",
            f"Desviación estándar: {numeric.std():.2f}",
        ])
        st.info(" | ".join(stats_lines))
    else:
        st.info(" | ".join(stats_lines))

    left, right = st.columns([1, 1.35])
    with left:
        st.markdown("### 🧾 Tabla de frecuencias")
        st.dataframe(table, use_container_width=True)
    with right:
        fig = build_univariable_chart(table, label, chart_type)
        st.plotly_chart(fig, use_container_width=True)

    interpretation = ""
    if mode_row is not None:
        interpretation = (
            f"La categoría con mayor presencia en **{label}** es **{mode_row['Categoría']}**, "
            f"con **{mode_row['Porcentaje']:.1f}%** de los registros válidos."
        )
        st.success(interpretation)

    if st.button("💾 Guardar este análisis completo en historial", type="primary"):
        add_history(
            f"Univariable | {label} | {st.session_state.active_file}",
            {
                "type": "Univariable",
                "label": label,
                "file": st.session_state.active_file,
                "table": table,
                "stats": "\n".join(stats_lines),
                "interpretation": interpretation,
                "chart_type": chart_type,
                "figure_json": fig.to_json(),
                "image_png": fig_to_png_bytes(fig),
            },
        )
        st.success("Análisis guardado con tabla, estadísticos e imagen/gráfico.")


def build_bivariable_chart(table: pd.DataFrame, pair_label: str, chart_type: str):
    if chart_type == "Heatmap":
        fig = px.imshow(
            table,
            text_auto=True,
            aspect="auto",
            title=f"Mapa de calor: {pair_label}",
            color_continuous_scale=["#EEF2FF", "#7C3AED", "#EC4899"],
        )
    else:
        plot_data = table.reset_index().melt(id_vars="Fila", var_name="Columna", value_name="Frecuencia")
        barmode = "group" if chart_type == "Barras agrupadas" else "stack"
        fig = px.bar(
            plot_data,
            x="Fila",
            y="Frecuencia",
            color="Columna",
            barmode=barmode,
            title=f"{chart_type}: {pair_label}",
            text="Frecuencia",
            color_discrete_sequence=CHART_COLORS,
        )
        fig.update_traces(textposition="outside")
    return style_figure(fig, fig.layout.title.text)


def bivariable_view() -> None:
    page_header("Análisis bivariable", "Cruza dos variables, calcula chi-cuadrado, p-valor y guarda gráficos completos.", "▦")
    data = active_data()
    if data is None:
        st.info("Carga un archivo Excel o CSV desde el menú lateral.")
        return

    st.markdown(f'<span class="history-pill">Archivo activo: {st.session_state.active_file}</span>', unsafe_allow_html=True)

    col_a, col_b = st.columns([1, 1])
    with col_a:
        mode = st.radio("Selección", ["Catálogo automático", "Columnas manuales"], horizontal=True)
    with col_b:
        chart_type = st.selectbox("Tipo de gráfico", ["Heatmap", "Barras agrupadas", "Barras apiladas"])

    if mode == "Catálogo automático":
        pair_label = st.selectbox("Cruce", [x[0] for x in BIVARIABLE_CATALOG])
        selected = next(x for x in BIVARIABLE_CATALOG if x[0] == pair_label)
        left_col = find_column(data, selected[1])
        right_col = find_column(data, selected[2])
        if left_col is None or right_col is None:
            st.error("No se encontraron una o ambas columnas para este cruce.")
            return
    else:
        c1, c2 = st.columns(2)
        with c1:
            left_col = st.selectbox("Variable fila", data.columns)
        with c2:
            right_col = st.selectbox("Variable columna", data.columns)
        pair_label = f"{left_col} vs {right_col}"

    left_values = as_labels(data[left_col]).reset_index(drop=True)
    right_values = as_labels(data[right_col]).reset_index(drop=True)
    row_count = min(len(left_values), len(right_values))
    temp = pd.DataFrame({"Fila": left_values.iloc[:row_count], "Columna": right_values.iloc[:row_count]}).dropna()
    temp = temp[(temp["Fila"] != "") & (temp["Columna"] != "")]

    table = pd.crosstab(temp["Fila"], temp["Columna"])
    if table.empty:
        st.warning("No hay pares válidos para analizar.")
        return

    chi2, p_value, dof, expected = chi2_contingency(table)

    section_title(pair_label)
    m1, m2, m3, m4 = st.columns(4)
    with m1:
        metric_card("Chi-cuadrado", f"{chi2:.3f}", "#7C3AED")
    with m2:
        metric_card("Grados de libertad", f"{int(dof)}", "#1A548C")
    with m3:
        metric_card("p-valor", f"{p_value:.4f}", "#F97316")
    with m4:
        metric_card("Pares válidos", f"{len(temp):,}", "#008C7A")

    if p_value < 0.05:
        interpretation = "Asociación significativa al 5%. Conviene revisar las celdas con mayor frecuencia para entender qué grupos explican la relación."
        st.success(interpretation)
    elif p_value < 0.10:
        interpretation = "Existe una señal moderada, aunque no concluyente al 5%. Puede ser útil aumentar la muestra o segmentar los datos."
        st.warning(interpretation)
    else:
        interpretation = "No hay evidencia suficiente de asociación estadística. Las diferencias podrían deberse a variación muestral."
        st.error(interpretation)

    left, right = st.columns([1.05, 1.35])
    with left:
        st.markdown("### 🧾 Tabla cruzada")
        st.dataframe(table, use_container_width=True)
    with right:
        fig = build_bivariable_chart(table, pair_label, chart_type)
        st.plotly_chart(fig, use_container_width=True)

    stats_text = f"Chi-cuadrado: {chi2:.3f}\ngl: {int(dof)}\np-valor: {p_value:.4f}\nPares válidos: {len(temp)}"

    if st.button("💾 Guardar este cruce completo en historial", type="primary"):
        table_to_save = table.reset_index()
        add_history(
            f"Bivariable | {pair_label} | {st.session_state.active_file} | p={p_value:.4f}",
            {
                "type": "Bivariable",
                "label": pair_label,
                "file": st.session_state.active_file,
                "table": table_to_save,
                "stats": stats_text,
                "interpretation": interpretation,
                "chart_type": chart_type,
                "figure_json": fig.to_json(),
                "image_png": fig_to_png_bytes(fig),
            },
        )
        st.success("Cruce guardado con tabla, estadísticos e imagen/gráfico.")


def build_excel_report() -> bytes:
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        workbook = writer.book
        title_fmt = workbook.add_format({"bold": True, "font_size": 16, "font_color": "#1A548C"})
        header_fmt = workbook.add_format({"bold": True, "bg_color": "#1A548C", "font_color": "#FFFFFF"})
        note_fmt = workbook.add_format({"text_wrap": True, "valign": "top"})

        summary_rows = []
        for item in st.session_state.history:
            summary_rows.append({"Historial": item})
        summary = pd.DataFrame(summary_rows if summary_rows else [{"Historial": "Sin historial"}])
        summary.to_excel(writer, index=False, sheet_name="Historial", startrow=3)
        ws = writer.sheets["Historial"]
        ws.write("A1", "Reporte de análisis de movilidad", title_fmt)
        ws.write("A2", f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", note_fmt)
        ws.set_column("A:A", 120)
        ws.set_row(3, None, header_fmt)

        for i, record in enumerate(st.session_state.records, start=1):
            sheet = f"{i:02d}_{record['type']}"[:31]
            record["table"].to_excel(writer, index=False, sheet_name=sheet, startrow=8)
            ws = writer.sheets[sheet]
            ws.write("A1", f"{record['type']}: {record['label']}", title_fmt)
            ws.write("A2", f"Archivo: {record.get('file', '')}", note_fmt)
            ws.write("A3", f"Fecha: {record.get('timestamp', '')}", note_fmt)
            ws.write("A4", "Estadísticos:", header_fmt)
            ws.write("A5", record.get("stats", ""), note_fmt)
            ws.write("B4", "Interpretación:", header_fmt)
            ws.write("B5", record.get("interpretation", ""), note_fmt)
            ws.set_column("A:A", 28)
            ws.set_column("B:Z", 20)
            ws.set_row(8, None, header_fmt)

            image_data = record.get("image_png")
            if not image_data:
                fig = fig_from_record(record)
                image_data = fig_to_png_bytes(fig) if fig is not None else None
            if image_data:
                ws.insert_image("H2", "grafico.png", {"image_data": io.BytesIO(image_data), "x_scale": 0.55, "y_scale": 0.55})
            else:
                ws.write("H2", "No se pudo incrustar la imagen. Instala kaleido para exportar gráficos como imagen.", note_fmt)
    return output.getvalue()


def build_docx_report() -> Optional[bytes]:
    if Document is None:
        return None

    document = Document()
    document.add_heading("Reporte de análisis de movilidad", level=1)
    document.add_paragraph(f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    document.add_heading("Historial general", level=2)
    for item in st.session_state.history:
        document.add_paragraph(item)

    for record in st.session_state.records:
        document.add_page_break()
        document.add_heading(f"{record['type']}: {record['label']}", level=2)
        document.add_paragraph(f"Archivo: {record.get('file', '')}")
        document.add_paragraph(f"Fecha: {record.get('timestamp', '')}")
        document.add_heading("Estadísticos", level=3)
        document.add_paragraph(record.get("stats", ""))
        document.add_heading("Interpretación", level=3)
        document.add_paragraph(record.get("interpretation", ""))

        image_data = record.get("image_png")
        if not image_data:
            fig = fig_from_record(record)
            image_data = fig_to_png_bytes(fig) if fig is not None else None
        if image_data and Inches is not None:
            document.add_heading("Gráfico", level=3)
            document.add_picture(io.BytesIO(image_data), width=Inches(6.3))
        else:
            document.add_paragraph("No se pudo incrustar el gráfico como imagen. Instala kaleido para habilitar esta opción.")

        document.add_heading("Tabla", level=3)
        table_df = record["table"].reset_index(drop=True)
        doc_table = document.add_table(rows=1, cols=len(table_df.columns))
        doc_table.style = "Table Grid"
        for j, col in enumerate(table_df.columns):
            doc_table.rows[0].cells[j].text = str(col)
        for _, row in table_df.head(120).iterrows():
            cells = doc_table.add_row().cells
            for j, value in enumerate(row):
                cells[j].text = str(value)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def history_view() -> None:
    page_header("Historial y exportación", "Aquí se guardan análisis completos: tablas, estadísticos, interpretaciones y gráficos.", "🕘")
    if not st.session_state.history:
        st.info("Aún no hay análisis guardados en esta sesión.")
        return

    section_title("Resumen de historial")
    st.text_area("Historial", "\n".join(st.session_state.history), height=180)

    section_title("Análisis guardados")
    for index, record in enumerate(st.session_state.records, start=1):
        with st.expander(f"{index}. {record['type']} | {record['label']} | {record.get('timestamp', '')}", expanded=index == len(st.session_state.records)):
            st.markdown(f'<span class="history-pill">{record.get("file", "Archivo no especificado")}</span>', unsafe_allow_html=True)
            c1, c2 = st.columns([1, 1.25])
            with c1:
                st.markdown("#### Estadísticos")
                st.code(record.get("stats", ""), language="text")
                st.markdown("#### Interpretación")
                st.write(record.get("interpretation", ""))
                st.markdown("#### Tabla")
                st.dataframe(record["table"], use_container_width=True)
            with c2:
                fig = fig_from_record(record)
                if fig is not None:
                    st.plotly_chart(
                        fig,
                        use_container_width=True,
                        key=f"history_plot_{index}_{record.get('timestamp', '')}_{record.get('type', '')}_{record.get('label', '')}",
                    )
                elif record.get("image_png"):
                    st.image(record["image_png"], caption="Gráfico guardado")
                else:
                    st.warning("Este registro no tiene gráfico disponible.")

    section_title("Descargar reportes")
    col1, col2 = st.columns(2)
    col1.download_button(
        "📥 Descargar reporte Excel con tablas e imágenes",
        build_excel_report(),
        "reporte_analisis_movilidad.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    docx = build_docx_report()
    if docx is not None:
        col2.download_button(
            "📥 Descargar reporte Word con tablas e imágenes",
            docx,
            "reporte_analisis_movilidad.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    else:
        col2.warning("Para exportar Word instala python-docx.")

    if st.button("🧹 Limpiar historial"):
        st.session_state.history = []
        st.session_state.records = []
        st.rerun()


def help_view() -> None:
    page_header("Ayuda", "Guía rápida para ejecutar y usar la app.", "📘")
    st.markdown(
        """
### Cómo correr la app
1. Instala dependencias: `pip install -r requirements.txt`
2. Ejecuta: `streamlit run app.py`
3. Abre la dirección local que muestra Streamlit.

### Archivos de datos
Puedes cargar `.xlsx`, `.xls` o `.csv` desde el menú lateral.

### Qué hace la app
- Permite cargar uno o varios archivos.
- Permite seleccionar el archivo activo.
- Ejecuta análisis univariable.
- Ejecuta análisis bivariable con chi-cuadrado, grados de libertad y p-valor.
- Genera gráficos interactivos con más color.
- Guarda en historial tablas, estadísticos, interpretación y gráficos.
- Exporta reportes en Excel y Word con tablas e imágenes.

### Nota para exportar imágenes
Para que Excel y Word incrusten los gráficos como imagen, instala también `kaleido`.
        """
    )


def main() -> None:
    init_state()
    inject_css()
    view = sidebar()
    if view == "Inicio":
        home_view()
    elif view == "Análisis univariable":
        univariable_view()
    elif view == "Análisis bivariable":
        bivariable_view()
    elif view == "Historial y exportación":
        history_view()
    else:
        help_view()


if __name__ == "__main__":
    main()
